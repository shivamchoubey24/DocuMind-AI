"""
DocuMind AI - Chat with multiple documents using RAG (Groq + FAISS + LangChain)
"""
import os
import time
from datetime import datetime

import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx

from langchain.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_groq import ChatGroq

from htmlTemplates import css, bot_template, user_template, source_template

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
VECTORSTORE_DIR = "vectorstore_index"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
SUPPORTED_EXTENSIONS = ("pdf", "txt", "docx")
FILE_ICONS = {"pdf": "📄", "txt": "📃", "docx": "📝"}

SUGGESTED_QUESTIONS = [
    "Summarize this document in 3 bullet points",
    "What are the key takeaways?",
    "Are there any dates, numbers, or figures worth noting?",
]

CUSTOM_QUESTION_PROMPT = PromptTemplate.from_template(
    """Given the following conversation and a follow up question, rephrase the
follow up question to be a standalone question, in its original language.

Chat History:
{chat_history}
Follow Up Input: {question}
Standalone question:"""
)


# ---------------------------------------------------------------------------
# Document loading (PDF / TXT / DOCX) -> LangChain Documents with metadata
# ---------------------------------------------------------------------------
def extract_documents(uploaded_files):
    """Extract text from uploaded files, keeping per-file / per-page metadata
    so answers can later be traced back to their source."""
    documents = []
    skipped = []

    for file in uploaded_files:
        filename = file.name
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        try:
            if ext == "pdf":
                pdf_reader = PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    text = page.extract_text() or ""
                    if text.strip():
                        documents.append(
                            Document(
                                page_content=text,
                                metadata={"source": filename, "page": page_num, "ext": ext},
                            )
                        )

            elif ext == "txt":
                text = file.read().decode("utf-8", errors="ignore")
                if text.strip():
                    documents.append(
                        Document(page_content=text, metadata={"source": filename, "page": 1, "ext": ext})
                    )

            elif ext == "docx":
                doc = docx.Document(file)
                text = "\n".join(p.text for p in doc.paragraphs)
                if text.strip():
                    documents.append(
                        Document(page_content=text, metadata={"source": filename, "page": 1, "ext": ext})
                    )
            else:
                skipped.append(filename)

        except Exception as e:
            st.error(f"⚠️ Couldn't process **{filename}**: {e}")

    if skipped:
        st.warning(
            "Skipped unsupported file(s): "
            + ", ".join(skipped)
            + f" — supported types are {', '.join(SUPPORTED_EXTENSIONS)}."
        )

    return documents


def get_chunks(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    return splitter.split_documents(documents)


def build_doc_library(chunks):
    """Summarize chunks per source file for the sidebar 'Document Library'."""
    library = {}
    for ch in chunks:
        src = ch.metadata.get("source", "unknown")
        ext = ch.metadata.get("ext", "")
        page = ch.metadata.get("page", 1)
        entry = library.setdefault(src, {"ext": ext, "chunks": 0, "pages": set()})
        entry["chunks"] += 1
        entry["pages"].add(page)
    return library


@st.cache_resource(show_spinner=False)
def _load_embeddings():
    # Cached so the (relatively slow) model load only happens once per session
    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL, model_kwargs={"device": "cpu"})


def get_vectorstore(chunks):
    embeddings = _load_embeddings()
    return FAISS.from_documents(documents=chunks, embedding=embeddings)


def save_vectorstore(vectorstore):
    vectorstore.save_local(VECTORSTORE_DIR)


def load_vectorstore():
    embeddings = _load_embeddings()
    return FAISS.load_local(
        VECTORSTORE_DIR, embeddings, allow_dangerous_deserialization=True
    )


def saved_index_exists():
    return os.path.isdir(VECTORSTORE_DIR) and os.listdir(VECTORSTORE_DIR)


def get_conversationchain(vectorstore):
    llm = ChatGroq(temperature=0.2, model_name="llama-3.1-8b-instant")
    memory = ConversationBufferMemory(
        memory_key="chat_history",
        return_messages=True,
        output_key="answer",
    )
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(search_kwargs={"k": 4}),
        condense_question_prompt=CUSTOM_QUESTION_PROMPT,
        memory=memory,
        return_source_documents=True,
    )
    return conversation_chain


# ---------------------------------------------------------------------------
# Chat handling
# ---------------------------------------------------------------------------
def handle_question(question):
    with st.spinner("Thinking..."):
        start = time.time()
        try:
            response = st.session_state.conversation({"question": question})
        except Exception as e:
            st.error(f"Something went wrong while generating a response: {e}")
            return
        elapsed = round(time.time() - start, 1)

    st.session_state.chat_history = response["chat_history"]
    st.session_state.sources.append(
        {"docs": response.get("source_documents", []), "elapsed": elapsed}
    )


def render_sources(entry):
    docs = entry["docs"]
    if not docs:
        return
    with st.expander(f"📄 {len(docs)} source(s) used — {entry['elapsed']}s to answer"):
        seen = set()
        rank = 0
        for doc in docs:
            src = doc.metadata.get("source", "unknown")
            page = doc.metadata.get("page", "?")
            ext = doc.metadata.get("ext", "")
            key = (src, page)
            if key in seen:
                continue
            seen.add(key)
            rank += 1
            icon = FILE_ICONS.get(ext, "📄")
            snippet = doc.page_content[:280].replace("\n", " ").strip()
            html = (
                source_template.replace("{{RANK}}", f"#{rank}")
                .replace("{{ICON}}", icon)
                .replace("{{SOURCE}}", src)
                .replace("{{PAGE}}", str(page))
                .replace("{{SNIPPET}}", snippet + ("..." if len(doc.page_content) > 280 else ""))
            )
            st.write(html, unsafe_allow_html=True)


def display_chat():
    if not st.session_state.chat_history:
        return
    for i, msg in enumerate(st.session_state.chat_history):
        if i % 2 == 0:
            st.write(user_template.replace("{{MSG}}", msg.content), unsafe_allow_html=True)
        else:
            st.write(bot_template.replace("{{MSG}}", msg.content), unsafe_allow_html=True)
            src_idx = i // 2
            if src_idx < len(st.session_state.sources):
                render_sources(st.session_state.sources[src_idx])


def build_transcript():
    lines = []
    if st.session_state.chat_history:
        for msg in st.session_state.chat_history:
            role = "You" if msg.type == "human" else "DocuMind AI"
            lines.append(f"{role}: {msg.content}")
    return "\n\n".join(lines)


def clear_chat_only():
    """Clear the visible conversation AND the chain's underlying memory,
    so old turns can't silently resurface on the next question."""
    if st.session_state.conversation is not None:
        st.session_state.conversation.memory.clear()
    st.session_state.chat_history = None
    st.session_state.sources = []


def reset_everything():
    st.session_state.conversation = None
    st.session_state.chat_history = None
    st.session_state.sources = []
    st.session_state.doc_stats = None
    st.session_state.doc_library = None


def ask(question_text):
    st.session_state.pending_question = question_text


# ---------------------------------------------------------------------------
# Main app
# ---------------------------------------------------------------------------
def main():
    load_dotenv()
    st.set_page_config(page_title="DocuMind AI - Chat with Documents", page_icon=":books:", layout="wide")
    st.write(css, unsafe_allow_html=True)

    if not os.getenv("GROQ_API_KEY"):
        st.error(
            "No `GROQ_API_KEY` found. Create a `.env` file in the project root "
            "with `GROQ_API_KEY=your-key-here` (see `.env.example`), then restart the app."
        )
        st.stop()

    # session state defaults
    st.session_state.setdefault("conversation", None)
    st.session_state.setdefault("chat_history", None)
    st.session_state.setdefault("sources", [])
    st.session_state.setdefault("doc_stats", None)
    st.session_state.setdefault("doc_library", None)
    st.session_state.setdefault("pending_question", "")

    st.markdown(
        '<div class="dm-header"><h1>📚 DocuMind AI</h1>'
        '<p>Chat with your PDF, TXT, and DOCX files — every answer cited back to its source.</p></div>',
        unsafe_allow_html=True,
    )

    has_convo = st.session_state.conversation is not None

    # If a suggestion chip was clicked, feed it in as the question for this run
    default_question = st.session_state.pending_question
    st.session_state.pending_question = ""

    question = st.text_input(
        "Ask a question about your documents:",
        value=default_question,
        disabled=not has_convo,
    )

    if not has_convo:
        st.info("👈 Upload documents and click **Process** in the sidebar to get started.")
    else:
        st.markdown('<div class="chip-row">', unsafe_allow_html=True)
        chip_cols = st.columns(len(SUGGESTED_QUESTIONS))
        for col, sq in zip(chip_cols, SUGGESTED_QUESTIONS):
            if col.button(sq, use_container_width=True, key=f"chip_{sq}"):
                ask(sq)
                st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    if question and has_convo:
        handle_question(question)

    display_chat()

    if st.session_state.chat_history:
        col1, col2 = st.columns([1, 1])
        with col1:
            st.download_button(
                "⬇️ Download chat transcript",
                data=build_transcript(),
                file_name=f"documind_chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with col2:
            if st.button("🗑️ Clear conversation", use_container_width=True):
                clear_chat_only()
                st.rerun()

    with st.sidebar:
        st.subheader("Your documents")
        docs = st.file_uploader(
            f"Upload files ({', '.join(SUPPORTED_EXTENSIONS)}) and click 'Process'",
            accept_multiple_files=True,
            type=list(SUPPORTED_EXTENSIONS),
        )

        col_a, col_b = st.columns(2)
        process_clicked = col_a.button("⚙️ Process", use_container_width=True)
        reset_clicked = col_b.button("🔄 Full reset", use_container_width=True)

        if reset_clicked:
            reset_everything()
            st.rerun()

        if process_clicked:
            if not docs:
                st.warning("Please upload at least one document first.")
            else:
                start = time.time()
                with st.spinner("Reading documents..."):
                    documents = extract_documents(docs)

                if not documents:
                    st.error("No readable text found in the uploaded file(s).")
                else:
                    with st.spinner("Splitting into chunks..."):
                        chunks = get_chunks(documents)

                    with st.spinner("Generating embeddings & building the vector store..."):
                        vectorstore = get_vectorstore(chunks)

                    with st.spinner("Setting up the conversation chain..."):
                        st.session_state.conversation = get_conversationchain(vectorstore)
                        st.session_state._last_vectorstore = vectorstore

                    elapsed = round(time.time() - start, 1)
                    st.session_state.doc_stats = {
                        "files": len({d.metadata["source"] for d in documents}),
                        "chunks": len(chunks),
                        "seconds": elapsed,
                    }
                    st.session_state.doc_library = build_doc_library(chunks)
                    st.session_state.chat_history = None
                    st.session_state.sources = []
                    st.success(f"Processed {len(docs)} file(s) in {elapsed}s. Ask away!")

        if st.session_state.doc_stats:
            stats = st.session_state.doc_stats
            s1, s2, s3 = st.columns(3)
            s1.metric("Files", stats["files"])
            s2.metric("Chunks", stats["chunks"])
            s3.metric("Time (s)", stats["seconds"])

        if st.session_state.doc_library:
            st.caption("📚 Document library")
            for fname, info in st.session_state.doc_library.items():
                icon = FILE_ICONS.get(info["ext"], "📄")
                page_label = f"{len(info['pages'])} page(s)" if info["ext"] == "pdf" else "text file"
                st.markdown(
                    f'<div class="doc-card"><span class="doc-name">{icon} {fname}</span>'
                    f'<span class="doc-meta">{info["chunks"]} chunks · {page_label}</span></div>',
                    unsafe_allow_html=True,
                )

        st.divider()
        st.caption("Persistent index")
        p1, p2 = st.columns(2)
        with p1:
            if st.button("💾 Save index", use_container_width=True, disabled=st.session_state.conversation is None):
                save_vectorstore(st.session_state._last_vectorstore)
                st.success("Index saved to disk.")
        with p2:
            if st.button("📂 Load saved", use_container_width=True, disabled=not saved_index_exists()):
                with st.spinner("Loading saved index..."):
                    vectorstore = load_vectorstore()
                    st.session_state.conversation = get_conversationchain(vectorstore)
                    st.session_state._last_vectorstore = vectorstore
                    st.session_state.chat_history = None
                    st.session_state.sources = []
                st.success("Loaded previously saved index. Ask away!")

        st.divider()
        with st.expander("ℹ️ How it works"):
            st.markdown(
                "1. **Extract** text from your PDF/TXT/DOCX files (page-aware for PDFs)\n"
                "2. **Chunk** it with overlap so context isn't lost\n"
                "3. **Embed** chunks locally with `all-MiniLM-L6-v2`\n"
                "4. **Store** vectors in FAISS for similarity search\n"
                "5. **Retrieve** the most relevant chunks per question\n"
                "6. **Generate** a grounded answer with Groq's Llama 3.1, with sources cited"
            )


if __name__ == "__main__":
    main()
